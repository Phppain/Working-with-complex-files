"""task3.py"""

from docx import Document
from docx.shared import Pt

def create_word_file():
    doc = Document()
    doc.add_paragraph('Hello Python')
    doc.save('hello_python.docx')

def read_bold_text():
    doc = Document('hello_python.docx')
    bold_text = []
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bold_text.append(run.text)
    return ' '.join(bold_text)

def create_formatted_word_file():
    doc = Document()
    p = doc.add_paragraph('This is a paragraph with custom formatting.')
    run = p.add_run(' This part has different formatting.')
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    doc.save('formatted_text.docx')

if __name__ == "__main__":
    create_word_file()
    print(read_bold_text())
    create_formatted_word_file()
